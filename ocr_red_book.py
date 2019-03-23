import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches

# map_size = (850, 60, 1980, 815)
map_size = (800, 125, 2100, 1000)
ready_img_size = (650, 430)

languages = 'rus+eng'  # используемые языки для распознания
prefix = 'temp/map_'  # временные картинки с картой
document = Document()

cur_pos = 8  # начальная и последующая позиция
constraint_pos = 182  # где заканчиваются насекомые

pattern_dict_insects = {'Класс': 'Отряд',
                        'Семейство': 'Категория и статус',
                        'Категория и статус': 'Описание',
                        'Места обитания и особенности биологии.': 'факторы'}

pattern_dict_animals = {'Класс': 'Отряд',
                        'Семейство': 'Категория и статус',
                        'Категория и статус': 'Описание',
                        'В Пензенск': 'Места обитания',
                        'Численность и лимитирующие': 'Источники информации'}


# По насек - семейство, вид, и распределить, где обитают
# Типа леса, поля, вода
# А по жив - сем, вид, статус, распростр по пенз обл, числ, меры охраны

def get_text(s, pattern):
    text = ''
    for start, end in pattern.items():
        if start == 'Класс' and s.find(start) == -1:
            text += s[string.find(string[0:3]):s.find(end)] + '\n'
            continue
        else:
            text += s[s.find(start):s.find(end)] + '\n'
    return text


def filter_content(s, it):
    if it < 89:
        cont = get_text(s, pattern_dict_insects)
    else:
        cont = get_text(s, pattern_dict_animals)

    content = [x.strip() for x in cont]
    for i, v in enumerate(content):
        if v == '':
            content[i] = ' '
    return ''.join(content)


def get_map(path, num):
    img = Image.open(path)
    map = img.crop(map_size)  # вырезаем нужную область
    map.thumbnail(ready_img_size, Image.ANTIALIAS)  # уменьшаем изображение
    map_path = prefix + str(num) + '.png'
    map.save(map_path)
    return map_path


while cur_pos != constraint_pos:
    local = 'pic_' + str(cur_pos) + '.png'
    page = Image.open(local)

    string = pytesseract.image_to_string(page, lang=languages)
    text = filter_content(string, cur_pos)

    cur_map_path = get_map(local, cur_pos)

    document.add_picture(cur_map_path, width=Inches(2.957717))
    document.add_paragraph(text=text)
    cur_pos += 1
    if cur_pos == 84:
        cur_pos += 5

document.save('red_book.docx')
